from flask import Flask, request, render_template, redirect, url_for
from docx import Document
import os

app = Flask(__name__)

def fill_template(template_path, replacements, output_dir):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    output_path = os.path.join(output_dir, f"{replacements['{name}']}.docx")
    doc.save(output_path)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        template_file = request.files['template']
        output_dir = request.form['output_dir']
        data = [{
            '{name}': name,
            '{company}': company,
            '{date_agreement}': date_agreement,
            '{sdm}': sdm,
            '{customer}': customer,
            '{start_date}': start_date,
            '{end_date}': end_date,
            '{hours}': hours,
            '{rate}': rate
        } for name, company, date_agreement, sdm, customer, start_date, end_date, hours, rate in zip(
            request.form.getlist('name[]'),
            request.form.getlist('company[]'),
            request.form.getlist('date_agreement[]'),
            request.form.getlist('sdm[]'),
            request.form.getlist('customer[]'),
            request.form.getlist('start_date[]'),
            request.form.getlist('end_date[]'),
            request.form.getlist('hours[]'),
            request.form.getlist('rate[]')
        )]

        # Save template to a temporary path
        template_path = os.path.join('temp', template_file.filename)
        os.makedirs('temp', exist_ok=True)
        template_file.save(template_path)

        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        for replacements in data:
            fill_template(template_path, replacements, output_dir)

        return f"Documents generated successfully in {output_dir}"
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
